# # import os
# # import sys

# # os.environ["TRANSFORMERS_OFFLINE"] = "1"
# # os.environ["HF_DATASETS_OFFLINE"] = "1"
# # os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# # from dotenv import load_dotenv
# # load_dotenv()

# # from fastapi import FastAPI, HTTPException
# # from fastapi.middleware.cors import CORSMiddleware
# # from pydantic import BaseModel
# # import sqlite3, hashlib

# # from langchain.chains import create_retrieval_chain
# # from langchain.chains.combine_documents import create_stuff_documents_chain
# # from langchain_core.prompts import ChatPromptTemplate
# # from langchain_groq import ChatGroq
# # from langchain_huggingface import HuggingFaceEmbeddings
# # from langchain_community.vectorstores import FAISS

# # # =========================================================
# # # APP
# # # =========================================================

# # app = FastAPI(title="MedicalAI Backend")

# # app.add_middleware(
# #     CORSMiddleware,
# #     allow_origins=["http://localhost:5173", "http://localhost:5174"],
# #     allow_credentials=True,
# #     allow_methods=["*"],
# #     allow_headers=["*"],
# # )

# # # =========================================================
# # # AUTH DATABASE
# # # =========================================================

# # DB = "users.db"

# # def init_db():
# #     conn = sqlite3.connect(DB)
# #     conn.execute("""CREATE TABLE IF NOT EXISTS users (
# #         id INTEGER PRIMARY KEY AUTOINCREMENT,
# #         name TEXT, email TEXT UNIQUE, password TEXT)""")
# #     conn.commit()
# #     conn.close()

# # init_db()

# # def hash_pw(pw):
# #     return hashlib.sha256(pw.encode()).hexdigest()

# # # =========================================================
# # # EMERGENCY DETECTION
# # # =========================================================

# # EMERGENCY_KEYWORDS = {
# #     "suicide": [
# #         "i want to die", "kill myself", "killing myself",
# #         "end my life", "suicide", "suicidal", "self harm",
# #         "self-harm", "hurt myself", "no reason to live",
# #         "don't want to live", "want to end it", "take my life",
# #         "better off dead", "can't go on", "give up on life",
# #         "rather be dead", "overdose on purpose",
# #     ],
# #     "heart_attack": [
# #         "chest pain", "chest pressure", "chest tightness",
# #         "left arm pain", "jaw pain", "heart attack",
# #         "shortness of breath", "heart is racing",
# #         "pain in chest", "crushing chest",
# #     ],
# #     "stroke": [
# #         "face drooping", "face numb", "slurred speech",
# #         "can't speak", "weakness on one side", "arm weakness",
# #         "sudden headache", "blurred vision", "stroke",
# #         "face falling",
# #     ],
# #     "severe_bleeding": [
# #         "bleeding badly", "won't stop bleeding", "blood everywhere",
# #         "severe bleeding", "cut myself badly", "bleeding out",
# #     ],
# #     "unconscious": [
# #         "not breathing", "unconscious", "passed out",
# #         "not waking up", "fainted", "collapsed",
# #         "no pulse", "not responding",
# #     ],
# # }

# # HELPLINE_MESSAGES = {
# #     "suicide": """🆘 YOU ARE NOT ALONE

# # It sounds like you are going through an incredibly difficult time. Your life has value and people care about you deeply.

# # 📞 PAKISTAN MENTAL HEALTH HELPLINES:
# # • Umang Helpline: 0317-4288665 (24/7)
# # • Rozan Counseling: 051-2890505
# # • Rescue Emergency: 1122
# # • Edhi Foundation: 115

# # Please reach out to someone you trust RIGHT NOW.""",

# #     "heart_attack": """🚨 POSSIBLE HEART ATTACK DETECTED

# # Your symptoms may indicate a MEDICAL EMERGENCY.

# # 📞 CALL IMMEDIATELY:
# # • Rescue / Ambulance: 1122
# # • Edhi Ambulance: 115
# # • Health Helpline: 1166

# # While waiting: Sit down, stay calm, chew aspirin if available. Do NOT drive yourself. Go to the nearest Emergency Room NOW.""",

# #     "stroke": """🚨 POSSIBLE STROKE DETECTED

# # Time is critical. Remember FAST:
# # • F - Face drooping?
# # • A - Arm weakness?
# # • S - Speech difficulty?
# # • T - Time to call 1122 NOW

# # 📞 CALL: 1122 (Rescue) or 115 (Edhi Ambulance)""",

# #     "severe_bleeding": """🚨 SEVERE BLEEDING EMERGENCY

# # 📞 CALL: 1122 (Rescue) or 115 (Edhi Ambulance)

# # Press firmly on the wound with a clean cloth. Keep pressure constant. Go to Emergency Room NOW.""",

# #     "unconscious": """🚨 UNCONSCIOUS PERSON — MEDICAL EMERGENCY

# # 📞 CALL: 1122 (Rescue) or 115 (Edhi Ambulance)

# # Check breathing. Start CPR if trained. Do not leave them alone.""",
# # }

# # def detect_emergency(text: str):
# #     text_lower = text.lower()
# #     for category, keywords in EMERGENCY_KEYWORDS.items():
# #         for kw in keywords:
# #             if kw in text_lower:
# #                 return category
# #     return None

# # # =========================================================
# # # RAG SETUP
# # # =========================================================

# # DB_FAISS_PATH = "vectorstore/db_faiss"

# # print("Loading embedding model...")
# # embedding_model = HuggingFaceEmbeddings(
# #     model_name="sentence-transformers/all-MiniLM-L6-v2",
# #     model_kwargs={"device": "cpu"},
# #     encode_kwargs={"normalize_embeddings": True}
# # )
# # print("Embedding model loaded")

# # print("Loading FAISS database...")
# # try:
# #     faiss_db = FAISS.load_local(
# #         DB_FAISS_PATH,
# #         embedding_model,
# #         allow_dangerous_deserialization=True
# #     )
# #     retriever = faiss_db.as_retriever(
# #         search_type="similarity",
# #         search_kwargs={"k": 5}
# #     )
# #     print("FAISS database loaded")
# # except Exception as e:
# #     print(f"ERROR loading FAISS: {e}")
# #     print(">>> Run 'python create_memory_for_llm.py' first to build the index!")
# #     faiss_db = None
# #     retriever = None

# # GROQ_API_KEY = os.getenv("GROQ_API_KEY")
# # if not GROQ_API_KEY:
# #     print("ERROR: GROQ_API_KEY not found in .env file!")

# # llm = ChatGroq(
# #     groq_api_key=GROQ_API_KEY,
# #     model_name="llama-3.1-8b-instant",
# #     temperature=0.5,
# #     max_tokens=512
# # )
# # print("LLM loaded")

# # prompt = ChatPromptTemplate.from_template("""
# # You are a helpful medical assistant with access to two knowledge sources:

# # 1. MEDICAL ENCYCLOPEDIA - Gale Encyclopedia of Medicine with detailed
# #    information about diseases, symptoms, treatments, and medications.

# # 2. LAHORE DOCTOR DIRECTORY - A list of specialist doctors in Lahore,
# #    Pakistan, including their names, specializations, clinics, and
# #    contact information.

# # Rules:
# # - If the user asks about a disease, symptom, treatment, or medication answer from the medical encyclopedia.
# # - If the user asks about a doctor, specialist, clinic, or hospital in Lahore answer from the doctor directory.
# # - If the user asks about both, answer both parts.
# # - If the answer is not in the context, say: "I don't have enough information in my documents to answer that."
# # - Never make up doctor names, phone numbers, or addresses.

# # Context:
# # {context}

# # Question:
# # {input}

# # Answer:
# # """)

# # if retriever:
# #     combine_docs_chain = create_stuff_documents_chain(llm, prompt)
# #     rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
# #     print("RAG chain ready!")
# # else:
# #     rag_chain = None

# # # =========================================================
# # # AUTH ROUTES
# # # =========================================================

# # class Register(BaseModel):
# #     name: str
# #     email: str
# #     password: str

# # class Login(BaseModel):
# #     email: str
# #     password: str

# # @app.post("/auth/register")
# # def register(r: Register):
# #     try:
# #         conn = sqlite3.connect(DB)
# #         conn.execute(
# #             "INSERT INTO users (name,email,password) VALUES (?,?,?)",
# #             (r.name, r.email, hash_pw(r.password))
# #         )
# #         conn.commit()
# #         conn.close()
# #         return {"message": "Registered successfully"}
# #     except Exception:
# #         raise HTTPException(status_code=400, detail="Email already exists")

# # @app.post("/auth/login")
# # def login(r: Login):
# #     conn = sqlite3.connect(DB)
# #     conn.row_factory = sqlite3.Row
# #     user = conn.execute(
# #         "SELECT * FROM users WHERE email=? AND password=?",
# #         (r.email, hash_pw(r.password))
# #     ).fetchone()
# #     conn.close()
# #     if not user:
# #         raise HTTPException(status_code=400, detail="Invalid email or password")
# #     return {
# #         "token": "tok_" + str(user["id"]),
# #         "name": user["name"],
# #         "email": user["email"],
# #         "uid": user["id"]
# #     }

# # @app.post("/auth/logout")
# # def logout():
# #     return {"message": "Logged out"}

# # ADMIN_EMAIL = "admin@medical.com"
# # ADMIN_PASSWORD = "admin123"

# # @app.post("/admin/login")
# # def admin_login(r: Login):
# #     if r.email != ADMIN_EMAIL or r.password != ADMIN_PASSWORD:
# #         raise HTTPException(status_code=401, detail="Invalid admin credentials")
# #     return {"token": "admin_tok", "role": "admin", "name": "Admin"}

# # @app.get("/admin/users")
# # def get_users():
# #     conn = sqlite3.connect(DB)
# #     conn.row_factory = sqlite3.Row
# #     users = conn.execute("SELECT id, name, email FROM users").fetchall()
# #     conn.close()
# #     return {"users": [dict(u) for u in users], "total": len(users)}

# # @app.delete("/admin/users/{user_id}")
# # def delete_user(user_id: int):
# #     conn = sqlite3.connect(DB)
# #     conn.execute("DELETE FROM users WHERE id=?", (user_id,))
# #     conn.commit()
# #     conn.close()
# #     return {"message": "User deleted"}

# # # ADMIN_EMAIL = "admin@medical.com"
# # # ADMIN_PASSWORD = "admin123"

# # # @app.post("/admin/login")
# # # def admin_login(r: Login):
# # #     if r.email != ADMIN_EMAIL or r.password != ADMIN_PASSWORD:
# # #         raise HTTPException(status_code=401, detail="Invalid admin credentials")
# # #     return {"token": "admin_tok", "role": "admin", "name": "Admin"}

# # # @app.get("/admin/users")
# # # def get_users():
# # #     conn = sqlite3.connect(DB)
# # #     conn.row_factory = sqlite3.Row
# # #     users = conn.execute("SELECT id, name, email FROM users").fetchall()
# # #     conn.close()
# # #     return {"users": [dict(u) for u in users], "total": len(users)}

# # # @app.delete("/admin/users/{user_id}")
# # # def delete_user(user_id: int):
# # #     conn = sqlite3.connect(DB)
# # #     conn.execute("DELETE FROM users WHERE id=?", (user_id,))
# # #     conn.commit()
# # #     conn.close()
# # #     return {"message": "User deleted"}
# # # =========================================================
# # # CHAT ROUTES
# # # =========================================================

# # class ChatMessage(BaseModel):
# #     message: str
# #     session_id: str = "default"

# # @app.post("/chat/send")
# # def chat_send(body: ChatMessage):
# #     user_message = body.message

# #     # Emergency check first
# #     emergency = detect_emergency(user_message)
# #     if emergency:
# #         return {"answer": HELPLINE_MESSAGES[emergency], "sources": []}

# #     if not rag_chain:
# #         raise HTTPException(
# #             status_code=503,
# #             detail="Vector store not loaded. Run 'python create_memory_for_llm.py' first."
# #         )

# #     try:
# #         response = rag_chain.invoke({"input": user_message})
# #         answer = response["answer"]
# #         answer += "\n\n⚠️ *This information is for educational purposes only and does not replace professional medical advice.*"

# #         sources = []
# #         seen = set()
# #         for doc in response.get("context", []):
# #             source = doc.metadata.get("source", "Unknown")
# #             page = doc.metadata.get("page", "?")
# #             source_type = doc.metadata.get("source_type", "unknown")
# #             key = f"{source}-{page}"
# #             if key not in seen:
# #                 seen.add(key)
# #                 sources.append({"source": source, "page": page, "type": source_type})

# #         return {"answer": answer, "sources": sources}

# #     except Exception as e:
# #         raise HTTPException(status_code=500, detail=str(e))

# # @app.post("/chat/save")
# # def chat_save(body: dict):
# #     return {"message": "Session saved"}

# # @app.get("/chat/history")
# # def chat_history():
# #     return []

# # @app.get("/")
# # def root():
# #     return {"status": "MedicalAI backend running"}
# import smtplib
# import secrets
# from email.mime.text import MIMEText
# import os
# os.environ["TRANSFORMERS_OFFLINE"] = "1"
# os.environ["HF_DATASETS_OFFLINE"] = "1"
# os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# from dotenv import load_dotenv
# load_dotenv()

# from docx import Document
# from fastapi import FastAPI, HTTPException
# from fastapi.responses import StreamingResponse
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# import sqlite3, hashlib, io

# from langchain.chains import create_retrieval_chain
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_groq import ChatGroq
# from langchain_huggingface import HuggingFaceEmbeddings
# from langchain_community.vectorstores import FAISS

# app = FastAPI(title="MedicalAI Backend")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173", "http://localhost:5174"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# DB = "users.db"

# def init_db():
#     conn = sqlite3.connect(DB)
#     conn.execute("""CREATE TABLE IF NOT EXISTS users (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         name TEXT, email TEXT UNIQUE, password TEXT)""")
#     conn.commit()
#     conn.close()

# init_db()

# def hash_pw(pw):
#     return hashlib.sha256(pw.encode()).hexdigest()

# EMERGENCY_KEYWORDS = {
#     "suicide": ["i want to die","kill myself","killing myself","end my life","suicide","suicidal","self harm","self-harm","hurt myself","no reason to live","don't want to live","want to end it","take my life","better off dead","can't go on","give up on life","rather be dead","overdose on purpose"],
#     "heart_attack": ["chest pain","chest pressure","chest tightness","left arm pain","jaw pain","heart attack","shortness of breath","heart is racing","pain in chest","crushing chest"],
#     "stroke": ["face drooping","face numb","slurred speech","can't speak","weakness on one side","arm weakness","sudden headache","blurred vision","stroke","face falling"],
#     "severe_bleeding": ["bleeding badly","won't stop bleeding","blood everywhere","severe bleeding","cut myself badly","bleeding out"],
#     "unconscious": ["not breathing","unconscious","passed out","not waking up","fainted","collapsed","no pulse","not responding"],
# }

# HELPLINE_MESSAGES = {
#     "suicide": "🆘 YOU ARE NOT ALONE\n\nPlease reach out immediately:\n• Umang Helpline: 0317-4288665 (24/7)\n• Rozan Counseling: 051-2890505\n• Rescue Emergency: 1122\n• Edhi Foundation: 115",
#     "heart_attack": "🚨 POSSIBLE HEART ATTACK\n\nCALL IMMEDIATELY:\n• Rescue: 1122\n• Edhi Ambulance: 115\n• Health Helpline: 1166\n\nSit down, stay calm, chew aspirin if available. Do NOT drive yourself.",
#     "stroke": "🚨 POSSIBLE STROKE\n\nRemember FAST:\n• F - Face drooping?\n• A - Arm weakness?\n• S - Speech difficulty?\n• T - Time to call 1122 NOW\n\nCALL: 1122 or 115",
#     "severe_bleeding": "🚨 SEVERE BLEEDING\n\nCALL: 1122 or 115\n\nPress firmly on wound with clean cloth. Go to Emergency Room NOW.",
#     "unconscious": "🚨 UNCONSCIOUS PERSON\n\nCALL: 1122 or 115\n\nCheck breathing. Start CPR if trained. Do not leave them alone.",
# }

# def detect_emergency(text):
#     text_lower = text.lower()
#     for category, keywords in EMERGENCY_KEYWORDS.items():
#         for kw in keywords:
#             if kw in text_lower:
#                 return category
#     return None

# DB_FAISS_PATH = "vectorstore/db_faiss"

# print("Loading embedding model...")
# embedding_model = HuggingFaceEmbeddings(
#     model_name="sentence-transformers/all-MiniLM-L6-v2",
#     model_kwargs={"device": "cpu"},
#     encode_kwargs={"normalize_embeddings": True}
# )
# print("Embedding model loaded")

# print("Loading FAISS database...")
# try:
#     faiss_db = FAISS.load_local(DB_FAISS_PATH, embedding_model, allow_dangerous_deserialization=True)
#     retriever = faiss_db.as_retriever(search_type="similarity", search_kwargs={"k": 5})
#     print("FAISS database loaded")
# except Exception as e:
#     print(f"ERROR loading FAISS: {e}")
#     faiss_db = None
#     retriever = None

# GROQ_API_KEY = os.getenv("GROQ_API_KEY")
# if not GROQ_API_KEY:
#     print("ERROR: GROQ_API_KEY not found in .env file!")

# llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name="llama-3.1-8b-instant", temperature=0.5, max_tokens=512)
# print("LLM loaded")

# prompt = ChatPromptTemplate.from_template("""
# You are a helpful medical assistant with access to two knowledge sources:
# 1. MEDICAL ENCYCLOPEDIA - Gale Encyclopedia of Medicine
# 2. LAHORE DOCTOR DIRECTORY - Specialist doctors in Lahore, Pakistan

# Rules:
# - Answer from the encyclopedia for diseases, symptoms, treatments, medications.
# - Answer from the doctor directory for doctors, specialists, clinics in Lahore.
# - If not in context, say: "I don't have enough information in my documents to answer that."
# - Never make up doctor names, phone numbers, or addresses.

# Context: {context}
# Question: {input}
# Answer:
# """)

# if retriever:
#     combine_docs_chain = create_stuff_documents_chain(llm, prompt)
#     rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
#     print("RAG chain ready!")
# else:
#     rag_chain = None

# class Register(BaseModel):
#     name: str
#     email: str
#     password: str

# class Login(BaseModel):
#     email: str
#     password: str

# class GoogleAuth(BaseModel):
#     email: str
#     name: str
#     google_id: str

# class ChatMessage(BaseModel):
#     message: str
#     session_id: str = "default"

# class DownloadRequest(BaseModel):
#     messages: list = []

# @app.post("/auth/register")
# def register(r: Register):
#     try:
#         conn = sqlite3.connect(DB)
#         conn.execute("INSERT INTO users (name,email,password) VALUES (?,?,?)", (r.name, r.email, hash_pw(r.password)))
#         conn.commit()
#         conn.close()
#         return {"message": "Registered successfully"}
#     except Exception:
#         raise HTTPException(status_code=400, detail="Email already exists")

# @app.post("/auth/login")
# def login(r: Login):
#     conn = sqlite3.connect(DB)
#     conn.row_factory = sqlite3.Row
#     user = conn.execute("SELECT * FROM users WHERE email=? AND password=?", (r.email, hash_pw(r.password))).fetchone()
#     conn.close()
#     if not user:
#         raise HTTPException(status_code=400, detail="Invalid email or password")
#     return {"token": "tok_" + str(user["id"]), "name": user["name"], "email": user["email"], "uid": user["id"]}

# @app.post("/auth/logout")
# def logout():
#     return {"message": "Logged out"}

# @app.post("/auth/google")
# def google_auth(r: GoogleAuth):
#     conn = sqlite3.connect(DB)
#     conn.row_factory = sqlite3.Row
#     user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
#     if not user:
#         conn.execute("INSERT INTO users (name,email,password) VALUES (?,?,?)", (r.name, r.email, hash_pw(r.google_id)))
#         conn.commit()
#         user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
#     conn.close()
#     return {"token": "tok_" + str(user["id"]), "name": user["name"], "email": user["email"], "uid": user["id"]}

# import smtplib
# import secrets
# from email.mime.text import MIMEText

# class ForgotPassword(BaseModel):
#     email: str

# @app.post("/auth/forgot-password")
# def forgot_password(r: ForgotPassword):
#     conn = sqlite3.connect(DB)
#     conn.row_factory = sqlite3.Row
#     user = conn.execute("SELECT * FROM users WHERE email=?", (r.email,)).fetchone()
#     conn.close()
#     if not user:
#         raise HTTPException(status_code=404, detail="Email not found")

#     token = secrets.token_urlsafe(32)
#     reset_link = f"http://localhost:5173/reset-password?token={token}&email={r.email}"

#     EMAIL_USER = os.getenv("EMAIL_USER")
#     EMAIL_PASS = os.getenv("EMAIL_PASS")

#     msg = MIMEText(f"Click this link to reset your password:\n\n{reset_link}\n\nThis link expires in 30 minutes.")
#     msg["Subject"] = "MedicalAI - Password Reset"
#     msg["From"] = EMAIL_USER
#     msg["To"] = r.email

#     try:
#         with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
#             server.login(EMAIL_USER, EMAIL_PASS)
#             server.send_message(msg)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Email failed: {str(e)}")

#     return {"message": "Reset email sent"}

# # =========================================================
# # ADMIN ROUTES
# # =========================================================

# ADMIN_EMAIL = "admin@medical.com"
# ADMIN_PASSWORD = "admin123"

# @app.post("/admin/login")
# def admin_login(r: Login):
#     if r.email != ADMIN_EMAIL or r.password != ADMIN_PASSWORD:
#         raise HTTPException(status_code=401, detail="Invalid admin credentials")
#     return {"token": "admin_tok", "role": "admin", "name": "Admin"}

# @app.get("/admin/users")
# def get_users():
#     conn = sqlite3.connect(DB)
#     conn.row_factory = sqlite3.Row
#     users = conn.execute("SELECT id, name, email FROM users").fetchall()
#     conn.close()
#     return {"users": [dict(u) for u in users], "total": len(users)}

# @app.delete("/admin/users/{user_id}")
# def delete_user(user_id: int):
#     conn = sqlite3.connect(DB)
#     conn.execute("DELETE FROM users WHERE id=?", (user_id,))
#     conn.commit()
#     conn.close()
#     return {"message": "User deleted"}

# # =========================================================
# # CHAT ROUTES
# # =========================================================

# @app.post("/chat/send")
# def chat_send(body: ChatMessage):
#     user_message = body.message
#     emergency = detect_emergency(user_message)
#     if emergency:
#         return {"answer": HELPLINE_MESSAGES[emergency], "sources": []}
#     if not rag_chain:
#         raise HTTPException(status_code=503, detail="Vector store not loaded.")
#     try:
#         response = rag_chain.invoke({"input": user_message})
#         answer = response["answer"]
#         answer += "\n\n⚠️ *This information is for educational purposes only and does not replace professional medical advice.*"
#         sources = []
#         seen = set()
#         for doc in response.get("context", []):
#             source = doc.metadata.get("source", "Unknown")
#             page = doc.metadata.get("page", "?")
#             key = f"{source}-{page}"
#             if key not in seen:
#                 seen.add(key)
#                 sources.append({"source": source, "page": page})
#         return {"answer": answer, "sources": sources}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/chat/save")
# def chat_save(body: dict):
#     return {"message": "Session saved"}

# @app.get("/chat/history")
# def chat_history():
#     return []

# @app.post("/chat/download/{session_id}")
# def download_chat(session_id: str, body: DownloadRequest):
#     doc = Document()
#     doc.add_heading("MedicalAI Chat Session", 0)
#     doc.add_paragraph(f"Session ID: {session_id}")
#     doc.add_paragraph("─" * 40)
#     for msg in body.messages:
#         role = "You" if msg.get("role") == "user" else "MedicalAI"
#         time = msg.get("time", "")
#         doc.add_paragraph(f"{role} [{time}]:")
#         doc.add_paragraph(msg.get("text", ""))
#         doc.add_paragraph("")
#     buf = io.BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return StreamingResponse(
#         buf,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         headers={"Content-Disposition": f"attachment; filename=MedicalAI_{session_id}.docx"}
#     )

# @app.get("/")
# def root():
#     return {"status": "MedicalAI backend running"}
import os
import smtplib
import secrets
import io
from email.mime.text import MIMEText

os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

from dotenv import load_dotenv
load_dotenv()

from docx import Document
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import hashlib
import psycopg2
from psycopg2.extras import RealDictCursor

from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

app = FastAPI(title="MedicalAI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================================================
# DATABASE SETUP (Neon PostgreSQL)
# =========================================================

DATABASE_URL = os.getenv("DATABASE_URL")

def get_conn():
    return psycopg2.connect(DATABASE_URL, sslmode="require")

def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        name TEXT,
        email TEXT UNIQUE,
        password TEXT
    )""")
    conn.commit()
    cur.close()
    conn.close()

init_db()

def hash_pw(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

# =========================================================
# EMERGENCY DETECTION
# =========================================================

EMERGENCY_KEYWORDS = {
    "suicide": ["i want to die","kill myself","killing myself","end my life","suicide","suicidal","self harm","self-harm","hurt myself","no reason to live","don't want to live","want to end it","take my life","better off dead","can't go on","give up on life","rather be dead","overdose on purpose"],
    "heart_attack": ["chest pain","chest pressure","chest tightness","left arm pain","jaw pain","heart attack","shortness of breath","heart is racing","pain in chest","crushing chest"],
    "stroke": ["face drooping","face numb","slurred speech","can't speak","weakness on one side","arm weakness","sudden headache","blurred vision","stroke","face falling"],
    "severe_bleeding": ["bleeding badly","won't stop bleeding","blood everywhere","severe bleeding","cut myself badly","bleeding out"],
    "unconscious": ["not breathing","unconscious","passed out","not waking up","fainted","collapsed","no pulse","not responding"],
}

HELPLINE_MESSAGES = {
    "suicide": "🆘 YOU ARE NOT ALONE\n\nPlease reach out immediately:\n• Umang Helpline: 0317-4288665 (24/7)\n• Rozan Counseling: 051-2890505\n• Rescue Emergency: 1122\n• Edhi Foundation: 115",
    "heart_attack": "🚨 POSSIBLE HEART ATTACK\n\nCALL IMMEDIATELY:\n• Rescue: 1122\n• Edhi Ambulance: 115\n• Health Helpline: 1166\n\nSit down, stay calm, chew aspirin if available. Do NOT drive yourself.",
    "stroke": "🚨 POSSIBLE STROKE\n\nRemember FAST:\n• F - Face drooping?\n• A - Arm weakness?\n• S - Speech difficulty?\n• T - Time to call 1122 NOW\n\nCALL: 1122 or 115",
    "severe_bleeding": "🚨 SEVERE BLEEDING\n\nCALL: 1122 or 115\n\nPress firmly on wound with clean cloth. Go to Emergency Room NOW.",
    "unconscious": "🚨 UNCONSCIOUS PERSON\n\nCALL: 1122 or 115\n\nCheck breathing. Start CPR if trained. Do not leave them alone.",
}

def detect_emergency(text):
    text_lower = text.lower()
    for category, keywords in EMERGENCY_KEYWORDS.items():
        for kw in keywords:
            if kw in text_lower:
                return category
    return None

# =========================================================
# RAG SETUP
# =========================================================

DB_FAISS_PATH = "vectorstore/db_faiss"

print("Loading embedding model...")
embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True}
)
print("Embedding model loaded")

print("Loading FAISS database...")
try:
    faiss_db = FAISS.load_local(DB_FAISS_PATH, embedding_model, allow_dangerous_deserialization=True)
    retriever = faiss_db.as_retriever(search_type="similarity", search_kwargs={"k": 5})
    print("FAISS database loaded")
except Exception as e:
    print(f"ERROR loading FAISS: {e}")
    faiss_db = None
    retriever = None

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name="llama-3.1-8b-instant", temperature=0.5, max_tokens=512)
print("LLM loaded")

prompt = ChatPromptTemplate.from_template("""
You are a helpful medical assistant with access to two knowledge sources:
1. MEDICAL ENCYCLOPEDIA - Gale Encyclopedia of Medicine
2. LAHORE DOCTOR DIRECTORY - Specialist doctors in Lahore, Pakistan

Rules:
- Answer from the encyclopedia for diseases, symptoms, treatments, medications.
- Answer from the doctor directory for doctors, specialists, clinics in Lahore.
- If not in context, say: "I don't have enough information in my documents to answer that."
- Never make up doctor names, phone numbers, or addresses.

Context: {context}
Question: {input}
Answer:
""")

if retriever:
    combine_docs_chain = create_stuff_documents_chain(llm, prompt)
    rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
    print("RAG chain ready!")
else:
    rag_chain = None

# =========================================================
# MODELS
# =========================================================

class Register(BaseModel):
    name: str
    email: str
    password: str

class Login(BaseModel):
    email: str
    password: str

class GoogleAuth(BaseModel):
    email: str
    name: str
    google_id: str

class ForgotPassword(BaseModel):
    email: str

class ChatMessage(BaseModel):
    message: str
    session_id: str = "default"

class DownloadRequest(BaseModel):
    messages: list = []

# =========================================================
# AUTH ROUTES
# =========================================================

@app.post("/auth/register")
def register(r: Register):
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("INSERT INTO users (name,email,password) VALUES (%s,%s,%s)", (r.name, r.email, hash_pw(r.password)))
        conn.commit()
        cur.close()
        conn.close()
        return {"message": "Registered successfully"}
    except Exception:
        raise HTTPException(status_code=400, detail="Email already exists")

@app.post("/auth/login")
def login(r: Login):
    conn = get_conn()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute("SELECT * FROM users WHERE email=%s AND password=%s", (r.email, hash_pw(r.password)))
    user = cur.fetchone()
    cur.close()
    conn.close()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid email or password")
    return {"token": "tok_" + str(user["id"]), "name": user["name"], "email": user["email"], "uid": user["id"]}

@app.post("/auth/logout")
def logout():
    return {"message": "Logged out"}

@app.post("/auth/google")
def google_auth(r: GoogleAuth):
    conn = get_conn()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute("SELECT * FROM users WHERE email=%s", (r.email,))
    user = cur.fetchone()
    if not user:
        cur.execute("INSERT INTO users (name,email,password) VALUES (%s,%s,%s)", (r.name, r.email, hash_pw(r.google_id)))
        conn.commit()
        cur.execute("SELECT * FROM users WHERE email=%s", (r.email,))
        user = cur.fetchone()
    cur.close()
    conn.close()
    return {"token": "tok_" + str(user["id"]), "name": user["name"], "email": user["email"], "uid": user["id"]}

@app.post("/auth/forgot-password")
def forgot_password(r: ForgotPassword):
    conn = get_conn()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute("SELECT * FROM users WHERE email=%s", (r.email,))
    user = cur.fetchone()
    cur.close()
    conn.close()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")

    reset_link = f"http://localhost:5173/reset-password?email={r.email}"
    EMAIL_USER = os.getenv("EMAIL_USER")
    EMAIL_PASS = os.getenv("EMAIL_PASS")

    msg = MIMEText(f"Hello {user['name']},\n\nClick this link to reset your password:\n\n{reset_link}\n\nIf you did not request this, ignore this email.")
    msg["Subject"] = "MedicalAI - Password Reset Request"
    msg["From"] = EMAIL_USER
    msg["To"] = r.email

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_USER, EMAIL_PASS)
            server.send_message(msg)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Email failed: {str(e)}")

    return {"message": "Reset email sent"}

# =========================================================
# ADMIN ROUTES
# =========================================================

ADMIN_EMAIL = "admin@medical.com"
ADMIN_PASSWORD = "admin123"

@app.post("/admin/login")
def admin_login(r: Login):
    if r.email != ADMIN_EMAIL or r.password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin credentials")
    return {"token": "admin_tok", "role": "admin", "name": "Admin"}

@app.get("/admin/users")
def get_users():
    conn = get_conn()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute("SELECT id, name, email FROM users")
    users = cur.fetchall()
    cur.close()
    conn.close()
    return {"users": [dict(u) for u in users], "total": len(users)}

@app.delete("/admin/users/{user_id}")
def delete_user(user_id: int):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM users WHERE id=%s", (user_id,))
    conn.commit()
    cur.close()
    conn.close()
    return {"message": "User deleted"}

# =========================================================
# CHAT ROUTES
# =========================================================

@app.post("/chat/send")
def chat_send(body: ChatMessage):
    user_message = body.message
    emergency = detect_emergency(user_message)
    if emergency:
        return {"answer": HELPLINE_MESSAGES[emergency], "sources": []}
    if not rag_chain:
        raise HTTPException(status_code=503, detail="Vector store not loaded.")
    try:
        response = rag_chain.invoke({"input": user_message})
        answer = response["answer"]
        answer += "\n\n⚠️ *This information is for educational purposes only and does not replace professional medical advice.*"
        sources = []
        seen = set()
        for doc in response.get("context", []):
            source = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            key = f"{source}-{page}"
            if key not in seen:
                seen.add(key)
                sources.append({"source": source, "page": page})
        return {"answer": answer, "sources": sources}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat/save")
def chat_save(body: dict):
    return {"message": "Session saved"}

@app.get("/chat/history")
def chat_history():
    return []

@app.post("/chat/download/{session_id}")
def download_chat(session_id: str, body: DownloadRequest):
    doc = Document()
    doc.add_heading("MedicalAI Chat Session", 0)
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph("─" * 40)
    for msg in body.messages:
        role = "You" if msg.get("role") == "user" else "MedicalAI"
        time = msg.get("time", "")
        doc.add_paragraph(f"{role} [{time}]:")
        doc.add_paragraph(msg.get("text", ""))
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=MedicalAI_{session_id}.docx"}
    )

@app.get("/")
def root():
    return {"status": "MedicalAI backend running"}